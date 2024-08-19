import os
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
import re
# 1

def generate_parameter(title, version, statement, date=""):
    """
    生成包含文档参数的字典，包括标题、作者、声明和日期。

    参数:
    title (str): 文档标题
    version (str): 版本号
    statement (str): 文档声明
    date (str): 文档日期，如果为空则使用当前日期

    返回:
    dict: 包含标题、作者、声明和日期的字典
    """
    # 如果没有提供日期，则使用当前日期
    if date == "":
        date = datetime.now().strftime("%Y-%m-%d")  # 获取当前日期并格式化为字符串

    # 创建包含文档参数的字典
    dir_parameter = {
        "title": title,
        "version": "版本号:" + version,
        "date": date,
        "statement": statement
    }

    # 返回包含文档参数的字典
    return dir_parameter

def generate_latex_document_pdf(left_header, right_header, cover_footer, urlid):
    """
    生成包含指定页眉和封面页脚内容的 LaTeX 文档模板，并将其保存到指定文件。

    参数:
    left_header (str): 页眉左侧内容
    right_header (str): 页眉右侧内容
    cover_footer (str): 封面页脚内容
    filename (str): 保存文件的路径
    """
    # LaTeX 文档模板 初版存放在内存中
    latex_template = f"""
\\usepackage{{needspace}}
\\usepackage{{fancyhdr}}
\\usepackage{{graphicx}}
\\usepackage{{amsmath}}
\\usepackage{{hyperref}}
\\usepackage{{geometry}}
\\usepackage{{xcolor}}
\\usepackage{{fontspec}}
\\usepackage{{tocloft}}
\\usepackage{{titlesec}}
\\usepackage{{longtable}}
\\usepackage{{booktabs}}
\\usepackage{{listings}}
\\usepackage{{xeCJK}}
\\usepackage{{etoolbox}}
\\usepackage{{array}}
\\usepackage{{caption}}
\\usepackage{{tabularx}}

% 页面布局
\\geometry{{
    a4paper,
    left=25mm,
    right=25mm,
    top=25mm,
    bottom=25mm,
}}

% 设置中文字体
\\setCJKmainfont{{SimSun}}  % 使用宋体作为中文主字体
\\setmainfont{{SimSun}}  % 设置英文字体为宋体

% 页眉页脚设置
\\fancypagestyle{{plain}}{{
    \\fancyhf{{}}
    \\fancyhead[L]{{{left_header}}}
    \\fancyhead[R]{{{right_header}}}
    \\fancyfoot[C]{{\\thepage}}  % 仅显示页码
}}

\\pagestyle{{plain}}

% 字体和颜色设置
\\colorlet{{mycolor}}{{blue}}
\\newcommand{{\\highlight}}[1]{{\\textbf{{\\textcolor{{mycolor}}{{#1}}}}}}

% 行间距设置
\\renewcommand{{\\baselinestretch}}{{1.5}}  % 调整行间距为1.5倍

% 目录设置
\\renewcommand{{\\contentsname}}{{\\centering 目录}}

% 调整目录的样式
\\renewcommand{{\\cftsecfont}}{{\\bfseries\\fontsize{{16pt}}{{16pt}}\\selectfont}}  % 三号字体
\\renewcommand{{\\cftsubsecfont}}{{\\bfseries\\fontsize{{15pt}}{{15pt}}\\selectfont}}  % 小三号字体
\\renewcommand{{\\cftsubsubsecfont}}{{\\bfseries\\fontsize{{15pt}}{{15pt}}\\selectfont}}  % 小三号字体
\\renewcommand{{\\cftsecpagefont}}{{\\bfseries}}
\\renewcommand{{\\cftsubsecpagefont}}{{\\bfseries}}
\\renewcommand{{\\cftsubsubsecpagefont}}{{\\bfseries}}
\\setlength{{\\cftbeforesecskip}}{{0.5em}}
\\setlength{{\\cftbeforesubsecskip}}{{0.2em}}

% 在每个一级标题前插入一个空行
\\pretocmd{{\\section}}{{\\addtocontents{{toc}}{{\\protect\\addvspace{{1.0\\baselineskip}}}}}}{{}}{{}}

% 目录设置为居中
\\renewcommand{{\\cfttoctitlefont}}{{\\hfill\\Huge\\bfseries}}
\\renewcommand{{\\cftaftertoctitle}}{{\\hfill}}

% 章节编号和标题格式
\\titleformat{{\\section}}{{\\normalfont\\Large\\bfseries}}{{\\thesection}}{{1em}}{{}}
\\titleformat{{\\subsection}}{{\\normalfont\\large\\bfseries}}{{\\thesubsection}}{{1em}}{{}}
\\titleformat{{\\subsubsection}}{{\\normalfont\\normalsize\\bfseries}}{{\\thesubsubsection}}{{1em}}{{}}

% 设置目录条目格式
\\setlength{{\\cftsecnumwidth}}{{3em}} % 设置章节编号的宽度
\\setlength{{\\cftsubsecnumwidth}}{{3.5em}} % 设置子章节编号的宽度
\\setlength{{\\cftsubsubsecnumwidth}}{{4em}} % 设置三级章节编号的宽度

% 强制显示章节和子章节编号
\\setcounter{{secnumdepth}}{{3}}
\\setcounter{{tocdepth}}{{3}}

% 调整目录各级标题的缩进
\\cftsetindents{{section}}{{1.5em}}{{3em}}
\\cftsetindents{{subsection}}{{3.5em}}{{3.5em}}
\\cftsetindents{{subsubsection}}{{7em}}{{4em}}

% 表格样式设置
\\captionsetup[table]{{skip=10pt}}
\\newcolumntype{{L}}[1]{{|>{{\\raggedright\\arraybackslash}}p{{#1}}|}}
\\newcolumntype{{C}}[1]{{|>{{\\centering\\arraybackslash}}p{{#1}}|}}
\\newcolumntype{{R}}[1]{{|>{{\\raggedleft\\arraybackslash}}p{{#1}}|}}

% 代码块设置
\\lstset{{
    basicstyle=\\ttfamily,
    breaklines=true,
    frame=single,
    backgroundcolor=\\color{{gray!10}},
    extendedchars=true,
    inputencoding=utf8,
    literate={{一}}{{\\CJKchar{{"4E00}}}}1
             {{二}}{{\\CJKchar{{"4E8C}}}}1
             {{三}}{{\\CJKchar{{"4E09}}}}1
             {{四}}{{\\CJKchar{{"56DB}}}}1
             {{五}}{{\\CJKchar{{"4E94}}}}1
             {{六}}{{\\CJKchar{{"516D}}}}1
             {{七}}{{\\CJKchar{{"4E03}}}}1
             {{八}}{{\\CJKchar{{"516B}}}}1
             {{九}}{{\\CJKchar{{"4E5D}}}}1
             {{零}}{{\\CJKchar{{"96F6}}}}1
}}

% 超链接设置
\\hypersetup{{
    colorlinks=true,
    linkcolor=black,  % 设置链接颜色为黑色
    urlcolor=black,   % 设置 URL 颜色为黑色
    filecolor=black,  % 设置文件链接颜色为黑色
    citecolor=black   % 设置引用颜色为黑色
}}

% 封面页面设置
\\newcommand{{\\coverpage}}[4]{{
    \\begin{{titlepage}}
        \\begin{{flushleft}}
            \\includegraphics[width=0.2\\textwidth]{{#4}}
        \\end{{flushleft}}
        \\centering
        \\vspace{{5cm}}
        {{\\Huge\\bfseries #1 \\par}}
        \\vspace{{1.5cm}}
        {{\\Large #2 \\par}}
        \\vspace{{1.5cm}}
        {{\\Large #3 \\par}}
        \\vfill
        {{\\Large {cover_footer} \\par}}
    \\end{{titlepage}}
}}

% 声明页面设置
\\newcommand{{\\statementpage}}[1]{{
    \\begin{{center}}
        \\vspace*{{2cm}}
        {{\\Large\\bfseries 声明 \\par}}
        \\vspace{{1.5cm}}
        {{\\large #1 \\par}}
        \\vfill
    \\end{{center}}
}}

% 图表标题设置
\\captionsetup[figure]{{
    labelformat=simple,
    labelsep=quad,
    font=small,
    justification=centering,
    format=hang,
    singlelinecheck=off
}}
\\renewcommand\\figurename{{图}}  % 设置图标题的前缀
\\renewcommand\\thefigure{{\\thesection.\\arabic{{figure}}}}  % 设置图编号格式为章节号.图号
\\makeatletter
\\@addtoreset{{figure}}{{section}}  % 在每个章节开始时重置图片编号
\\makeatother

    """

    filename = os.path.join('trans_docx', urlid, "document_pdf.tex")
    # filename = os.path.join(urlid, "document_pdf.tex")

    # 检查文件路径的目录是否存在，如果不存在则创建目录
    os.makedirs(os.path.dirname(filename), exist_ok=True)

    # 打开文件进行写入，如果文件不存在则创建文件
    with open(filename, "w", encoding="utf-8") as file:
        file.write(latex_template)
    print(f"File '{filename}' has been created/overwritten with the provided content.")

    return filename


# 默认模板
def generate_latex_document_no_header_footer(urlid):
    """
    生成一个没有页眉和页脚的简单 LaTeX 文档模板，并将其保存到指定文件。

    参数:
    urlid (str): 文件路径的唯一标识符
    """
    # LaTeX 文档模板存放在内存中
    latex_template = f"""
\\usepackage{{graphicx}}
\\usepackage{{amsmath}}
\\usepackage{{hyperref}}
\\usepackage{{geometry}}
\\usepackage{{xcolor}}
\\usepackage{{fontspec}}
\\usepackage{{titlesec}}
\\usepackage{{longtable}}
\\usepackage{{booktabs}}
\\usepackage{{listings}}
\\usepackage{{xeCJK}}
\\usepackage{{etoolbox}}
\\usepackage{{array}}
\\usepackage{{caption}}
\\usepackage{{tabularx}}
\\usepackage{{needspace}}  % 添加对 needspace 包的引用

% 页面布局
\\geometry{{
    a4paper,
    left=25mm,
    right=25mm,
    top=25mm,
    bottom=25mm,
}}

% 设置中文字体
\\setCJKmainfont{{SimSun}}  % 使用宋体作为中文主字体
\\setmainfont{{SimSun}}  % 设置英文字体为宋体

% 取消页眉和页脚
\\pagestyle{{empty}}

% 字体和颜色设置
\\colorlet{{mycolor}}{{blue}}
\\newcommand{{\\highlight}}[1]{{\\textbf{{\\textcolor{{mycolor}}{{#1}}}}}}

% 行间距设置
\\renewcommand{{\\baselinestretch}}{{1.5}}  % 调整行间距为1.5倍

% 章节编号和标题格式
\\titleformat{{\\section}}{{\\normalfont\\Large\\bfseries}}{{\\thesection}}{{1em}}{{}}
\\titleformat{{\\subsection}}{{\\normalfont\\large\\bfseries}}{{\\thesubsection}}{{1em}}{{}}
\\titleformat{{\\subsubsection}}{{\\normalfont\\normalsize\\bfseries}}{{\\thesubsubsection}}{{1em}}{{}}

% 强制显示章节和子章节编号
\\setcounter{{secnumdepth}}{{3}}
\\setcounter{{tocdepth}}{{3}}

% 表格样式设置
\\captionsetup[table]{{skip=10pt}}
\\newcolumntype{{L}}[1]{{|>{{\\raggedright\\arraybackslash}}p{{#1}}|}}
\\newcolumntype{{C}}[1]{{|>{{\\centering\\arraybackslash}}p{{#1}}|}}
\\newcolumntype{{R}}[1]{{|>{{\\raggedleft\\arraybackslash}}p{{#1}}|}}

% 代码块设置
\\lstset{{
    basicstyle=\\ttfamily,
    breaklines=true,
    frame=single,
    backgroundcolor=\\color{{gray!10}},
    extendedchars=true,
    inputencoding=utf8,
    literate={{一}}{{\\CJKchar{{"4E00}}}}1
             {{二}}{{\\CJKchar{{"4E8C}}}}1
             {{三}}{{\\CJKchar{{"4E09}}}}1
             {{四}}{{\\CJKchar{{"56DB}}}}1
             {{五}}{{\\CJKchar{{"4E94}}}}1
             {{六}}{{\\CJKchar{{"516D}}}}1
             {{七}}{{\\CJKchar{{"4E03}}}}1
             {{八}}{{\\CJKchar{{"516B}}}}1
             {{九}}{{\\CJKchar{{"4E5D}}}}1
             {{零}}{{\\CJKchar{{"96F6}}}}1
}}

% 超链接设置
\\hypersetup{{
    colorlinks=true,
    linkcolor=black,  % 设置链接颜色为黑色
    urlcolor=black,   % 设置 URL 颜色为黑色
    filecolor=black,  % 设置文件链接颜色为黑色
    citecolor=black   % 设置引用颜色为黑色
}}

% 图表标题设置
\\captionsetup[figure]{{
    labelformat=simple,
    labelsep=quad,
    font=small,
    justification=centering,
    format=hang,
    singlelinecheck=off
}}
\\renewcommand\\figurename{{图}}  % 设置图标题的前缀
\\renewcommand\\thefigure{{\\thesection.\\arabic{{figure}}}}  % 设置图编号格式为章节号.图号
\\makeatletter
\\@addtoreset{{figure}}{{section}}  % 在每个章节开始时重置图片编号
\\makeatother

    """

    filename = os.path.join('trans_docx', urlid, "document_no.tex")

    # 检查文件路径的目录是否存在，如果不存在则创建目录
    os.makedirs(os.path.dirname(filename), exist_ok=True)

    # 打开文件进行写入，如果文件不存在则创建文件
    with open(filename, "w", encoding="utf-8") as file:
        file.write(latex_template)
    print(f"File '{filename}' has been created/overwritten with the provided content.")

    return filename



# 创建带有页眉的模板
def create_template_with_headers(template_path, left_header, right_header):
    """
    创建一个包含页眉和页脚的DOCX模板。

    参数:
        template_path (str): 模板文件保存路径。
        left_header (str): 左页眉内容。
        right_header (str): 右页眉内容。
    """
    # 在路径前添加 trans_docx/
    template_path = os.path.join('trans_docx', template_path)

    # 创建新的文档
    doc = Document()

    # 创建页眉
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = left_header
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 添加制表符用于分隔左右页眉
    tab_stops = header_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    header_paragraph.add_run('\t\t')  # 制表符用于分隔左右页眉
    header_paragraph.add_run(right_header)

    # 创建页脚并添加页码
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.add_run("Page ")
    fldChar1 = OxmlElement('w:fldChar')  # 开始页码字段
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE   \\* MERGEFORMAT"  # 添加PAGE命令
    fldChar2 = OxmlElement('w:fldChar')  # 结束页码字段
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # 保存模板
    doc.save(template_path)
    print(f"Template with cover and headers created at {template_path}")


# 添加封面页
def add_cover_page(doc, title, version, date, statement):
    """
    向DOCX文档中添加封面页。

    参数:
        doc (Document): DOCX文档对象。
        title (str): 文档标题。
        version (str): 版本号。
        date (str): 文档日期。
        statement (str): 可选声明。
    """
    # 添加封面页
    section = doc.sections[0]

    # 标题
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 空行
    doc.add_paragraph()

    # 作者
    version_paragraph = doc.add_paragraph()
    version_run = version_paragraph.add_run(version)
    version_run.font.size = Pt(18)
    version_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 空行
    doc.add_paragraph()

    # 时间
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(date)
    date_run.font.size = Pt(18)
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加声明页
    if statement:
        # 添加分页符
        doc.add_page_break()

        # 添加一个段落用于显示“声明”
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run("声明")
        title_run.font.size = Pt(14)  # 可以调整字体大小
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加一个段落用于显示声明内容
        statement_paragraph = doc.add_paragraph()
        statement_run = statement_paragraph.add_run(statement)
        statement_run.font.size = Pt(12)
        statement_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# 添加目录页
def add_table_of_contents(doc):
    """
    向DOCX文档中添加目录页。

    参数:
        doc (Document): DOCX文档对象。
    """
    toc_paragraph = doc.add_paragraph()
    run = toc_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # 添加分页符
    doc.add_page_break()


def add_header_image_to_first_page(doc, image_path, right_text):
    """
    为文档的第一页添加左页眉图片。

    :param doc: Document对象
    :param image_path: str, 要添加的图片的路径。
    :param right_text: str, 要添加的右页眉的文本。
    """
    # 在路径前添加 trans_docx/
    image_path = os.path.join('trans_docx', image_path)

    section = doc.sections[0]

    # 确保首页页眉不同
    section.header.is_linked_to_previous = False
    section.different_first_page_header_footer = True

    # 获取首页页眉
    header = section.first_page_header
    # 清空现有的页眉内容
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._element = None

    # # 添加图片到首页页眉
    # header_paragraph = header.add_paragraph()
    # header_paragraph.alignment = 0  # 左对齐
    # run = header_paragraph.add_run()
    # run.add_picture(image_path, width=Inches(1.0))  # 调整宽度，单位是英寸

        # 添加图片到首页页眉左侧
    left_paragraph = header.add_paragraph()
    left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = left_paragraph.add_run()
    run.add_picture(image_path, width=Inches(1.0))  # 调整宽度，单位是英寸

    # 添加文本到首页页眉右侧
    right_paragraph = header.add_paragraph()
    right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = right_paragraph.add_run(right_text)
    run.font.size = Pt(12)  # 设置字体大小

    return doc


# 更新目录
def update_toc(docx_file_path):
    """
    更新DOCX文档中的目录。

    参数:
        docx_file_path (str): DOCX文档路径。
    """
    # 在路径前添加 trans_docx/
    docx_file_path = os.path.join('trans_docx', docx_file_path)
    # 打开文档
    doc = Document(docx_file_path)

    # 找到目录并更新它
    for paragraph in doc.paragraphs:
        if 'TOC \\o "1-3"' in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            break

    # 保存文档
    doc.save(docx_file_path)


# 应用页眉和页脚到所有章节
def apply_headers_footers_to_sections(doc, left_header, right_header):
    """
    向DOCX文档中的所有章节应用页眉和页脚。

    参数:
        doc (Document): DOCX文档对象。
        left_header (str): 左页眉内容。
        right_header (str): 右页眉内容。
    """
    for section in doc.sections:
        # 页眉
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = left_header
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        tab_stops = header_paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        header_paragraph.add_run('\t')
        header_paragraph.add_run(right_header)

        # 页脚
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 清除可能存在的重复页码字段
        for run in footer_paragraph.runs:
            run.clear()

        run = footer_paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
