import os
import subprocess
from util.generate import add_cover_page\
    , add_table_of_contents, update_toc\
    , apply_headers_footers_to_sections\
    , add_header_image_to_first_page
from docx import Document
from docxcompose.composer import Composer
from PIL import Image


def get_image_height(image_path):
    """获取图片的高度（单位为 pt，1 pt = 1/72 英寸）"""
    with Image.open(image_path) as img:
        # 获取图片的高度（像素）
        height_px = img.height
        # 假设图像的分辨率为 72 DPI（默认），将高度转换为点（pt）
        height_pt = height_px * 72 / img.info.get('dpi', (72, 72))[1]
        # print(height_pt)
        
        return height_pt
    

# md -> pdf
def convert_markdown_to_pdf_use(urlid, input_file, title, version, date,
                             output_file, logo_path, header_file, resource_paths=[],
                            statement=""):
    """
    将Markdown文件转换为PDF文件。

    参数:
        input_file (str): 输入的Markdown文件路径。
        title (str): 文档标题。
        version (str): 版本号。
        date (str): 文档日期。
        output_file (str): 输出的PDF文件路径。
        header_file (str): LaTeX header文件路径。
        logo_path (str): logo文件路径。
        resource_paths (list): 资源文件路径列表。
        statement (str): 可选声明。
    """
    # 将路径标准化并替换反斜杠为正斜杠
    input_file = input_file.replace("\\", "/")
    output_file = output_file.replace("\\", "/")
    header_file = header_file.replace("\\", "/")
    logo_path = logo_path.replace("\\", "/")
    resource_paths = [path.replace("\\", "/") for path in resource_paths]

    # # 将资源路径列表转换为字符串，使用冒号分隔
    # resource_path_str = ":".join(resource_paths)
    # 将资源路径列表转换为字符串，使用操作系统的路径分隔符
    # print("=========================resource_paths===========================", resource_paths)
    resource_path_str = os.pathsep.join(resource_paths)
    # print("=========================resource_path_str===========================", resource_path_str)

    # 创建一个临时的Markdown文件，用于存储转换过程中的中间数据
    temp_md_file = os.path.join(os.path.dirname(input_file), "temp.md")
    with open(temp_md_file, "w", encoding="utf-8") as f:
        # 写入封面信息，包含标题、作者、日期和logo
        f.write(f"\\coverpage{{{title}}}{{{version}}}{{{date}}}{{{logo_path}}}\n\n")
        f.write("\\newpage\n\n")

        # 如果有声明信息，则写入声明信息
        if statement:
            f.write(f"\\statementpage{{{statement}}}\n\n")
            f.write("\\newpage\n\n")

        # 写入目录页
        f.write("\\tableofcontents\n\n")
        f.write("\\newpage\n\n")

        # 读取原始Markdown文件内容，并写入临时Markdown文件
        with open(input_file, "r", encoding="utf-8") as original_md:
            previous_line = ""
            for line in original_md:
                # 每次遇到Markdown标题时在前面添加空行
                if line.strip().startswith("#"):
                    if previous_line.strip():
                        f.write("\n")
                    f.write("\n" + line.strip() + "\n\n")
                # 如果检测到图片标签，则确保图片上下有空行
                elif line.strip().startswith("![") and "](" in line:
                    if previous_line.strip():
                        f.write("\n")
                    
                    # 提取图片路径
                    start_idx = line.find("](") + 2
                    end_idx = line.find(")", start_idx)
                    image_path = line[start_idx:end_idx]

                    # print(image_path)
                    if image_path.split("/")[0] == ".":
                        image_path_str = image_path.split("/", 1)
                        # print(image_path_str)
                        image_path_end = image_path_str[0] + "/" + "trans_docx" + "/" + urlid + "/" + image_path_str[1]
                        # image_path_end = image_path_str[0] + "/" + urlid + "/" + image_path_str[1]
                        
                        print(image_path_end)
                    else:
                        image_path_end = "trans_docx/" + urlid + "/" + image_path
                        # image_path_end = urlid + "/" + image_path


                    # 获取图片的高度并插入\needspace命令
                    image_height = get_image_height(image_path_end)
                    # print(image_height)
                    f.write(f"\\needspace{{{image_height + 10}pt}}\n")
                    f.write(line.strip() + "\n\n")
                else:
                    f.write(line)
                previous_line = line

    # # 打印资源路径字符串，供调试使用
    # print(resource_path_str)

    # Pandoc命令，用于将Markdown转换为PDF
    command = [
        "pandoc",
        temp_md_file,  # 输入文件为临时Markdown文件
        "-o", output_file,  # 输出文件为指定的PDF文件
        "--pdf-engine=xelatex",  # 使用xelatex引擎
        f"--include-in-header={header_file}",  # 包含指定的LaTeX header文件
        "--resource-path", resource_path_str,  # 资源路径
        "-V", "tables=true",  # 启用表格支持
        "-V", "longtable=true",  # 启用长表格支持
        "-V", "booktabs=true",  # 启用booktabs支持
        "--listings",  # 启用代码高亮
        "--highlight-style=pygments",  # 使用pygments代码高亮样式
        "-V", "geometry:margin=1in",  # 设置页面边距
    ]

    # 运行Pandoc命令
    result = subprocess.run(command, cwd=os.path.dirname(input_file), capture_output=True, text=True)

    # 检查命令执行结果，如果出错则打印错误信息
    if result.returncode != 0:
        print(f"Error converting {input_file} to {output_file}")
        print(result.stderr)

    # 删除临时Markdown文件
    os.remove(temp_md_file)


def convert_markdown_to_pdf_basic(urlid, input_file,
                             output_file,header_file, resource_paths=[],
                            ):
    """
    将Markdown文件转换为PDF文件。

    参数:
        input_file (str): 输入的Markdown文件路径。
        title (str): 文档标题。
        version (str): 版本号。
        date (str): 文档日期。
        output_file (str): 输出的PDF文件路径。
        header_file (str): LaTeX header文件路径。
        logo_path (str): logo文件路径。
        resource_paths (list): 资源文件路径列表。
        statement (str): 可选声明。
    """
    # 将路径标准化并替换反斜杠为正斜杠
    input_file = input_file.replace("\\", "/")
    output_file = output_file.replace("\\", "/")
    header_file = header_file.replace("\\", "/")
    resource_paths = [path.replace("\\", "/") for path in resource_paths]

    # # 将资源路径列表转换为字符串，使用冒号分隔
    # resource_path_str = ":".join(resource_paths)
    # 将资源路径列表转换为字符串，使用操作系统的路径分隔符
    # print("=========================resource_paths===========================", resource_paths)
    resource_path_str = os.pathsep.join(resource_paths)
    # print("=========================resource_path_str===========================", resource_path_str)

    # 创建一个临时的Markdown文件，用于存储转换过程中的中间数据
    temp_md_file = os.path.join(os.path.dirname(input_file), "temp.md")
    with open(temp_md_file, "w", encoding="utf-8") as f:

        # 读取原始Markdown文件内容，并写入临时Markdown文件
        with open(input_file, "r", encoding="utf-8") as original_md:
            previous_line = ""
            for line in original_md:
                # 每次遇到Markdown标题时在前面添加空行
                if line.strip().startswith("#"):
                    if previous_line.strip():
                        f.write("\n")
                    f.write("\n" + line.strip() + "\n\n")
                # 如果检测到图片标签，则确保图片上下有空行
                elif line.strip().startswith("![") and "](" in line:
                    if previous_line.strip():
                        f.write("\n")
                    
                    # 提取图片路径
                    start_idx = line.find("](") + 2
                    end_idx = line.find(")", start_idx)
                    image_path = line[start_idx:end_idx]

                    # print(image_path)
                    if image_path.split("/")[0] == ".":
                        image_path_str = image_path.split("/", 1)
                        # print(image_path_str)
                        image_path_end = image_path_str[0] + "/" + "trans_docx" + "/" + urlid + "/" + image_path_str[1]
                        # image_path_end = image_path_str[0] + "/" + urlid + "/" + image_path_str[1]
                        
                        print(image_path_end)
                    else:
                        image_path_end = "trans_docx/" + urlid + "/" + image_path
                        # image_path_end = urlid + "/" + image_path


                    # 获取图片的高度并插入\needspace命令
                    image_height = get_image_height(image_path_end)
                    # print(image_height)
                    f.write(f"\\needspace{{{image_height + 10}pt}}\n")
                    f.write(line.strip() + "\n\n")
                else:
                    f.write(line)
                previous_line = line

    # # 打印资源路径字符串，供调试使用
    # print(resource_path_str)

    # Pandoc命令，用于将Markdown转换为PDF
    command = [
        "pandoc",
        temp_md_file,  # 输入文件为临时Markdown文件
        "-o", output_file,  # 输出文件为指定的PDF文件
        "--pdf-engine=xelatex",  # 使用xelatex引擎
        f"--include-in-header={header_file}",  # 包含指定的LaTeX header文件
        "--resource-path", resource_path_str,  # 资源路径
        "-V", "tables=true",  # 启用表格支持
        "-V", "longtable=true",  # 启用长表格支持
        "-V", "booktabs=true",  # 启用booktabs支持
        "--listings",  # 启用代码高亮
        "--highlight-style=pygments",  # 使用pygments代码高亮样式
        "-V", "geometry:margin=1in",  # 设置页面边距
    ]

    # 运行Pandoc命令
    result = subprocess.run(command, cwd=os.path.dirname(input_file), capture_output=True, text=True)

    # 检查命令执行结果，如果出错则打印错误信息
    if result.returncode != 0:
        print(f"Error converting {input_file} to {output_file}")
        print(result.stderr)

    # 删除临时Markdown文件
    os.remove(temp_md_file)



# md -> html
def convert_markdown_to_html(input_file, output_file, resource_paths=[], title="Document"):
    """
    将Markdown文件转换为HTML文件。

    参数:
        input_file (str): 输入的Markdown文件路径。
        output_file (str): 输出的HTML文件路径。
        resource_paths (list): 资源文件路径列表。
        title (str): 文档标题。
    """
    # # 将资源路径列表转换为字符串，使用冒号分隔
    # resource_path_str = ":".join(resource_paths)
    # 将资源路径列表转换为字符串，使用操作系统的路径分隔符
    resource_path_str = os.pathsep.join(resource_paths)
    print(resource_path_str)

    # 创建一个临时的Markdown文件，用于存储转换过程中的中间数据
    temp_md_file = os.path.join(os.path.dirname(input_file), "temp.md")

    # 确保styles.css文件存在，如果不存在则创建一个默认的styles.css文件
    css_path = os.path.join(os.getcwd(), "templates/styles.css")
    # print(css_path)
    if not os.path.exists(css_path):
        # 创建并写入改进后的CSS样式
        with open(css_path, "w", encoding="utf-8") as f:
            f.write("""
/* 设置整体页面的字体和背景颜色 */
body {
    font-family: Arial, sans-serif;  /* 使用Arial字体，后备字体为sans-serif */
    margin: 20px;  /* 设置页面边距 */
    background-color: #f9f9f9;  /* 设置背景颜色 */
    color: #333;  /* 设置文本颜色 */
}

/* 设置各级标题的颜色和间距 */
h1, h2, h3, h4, h5, h6 {
    color: #444;  /* 设置标题颜色 */
    margin-top: 1.2em;  /* 设置标题顶部间距 */
    margin-bottom: 0.6em;  /* 设置标题底部间距 */
}

/* 设置一级标题的字体大小和底部边框 */
h1 {
    font-size: 2.5em;  /* 设置字体大小 */
    border-bottom: 2px solid #ddd;  /* 设置底部边框 */
    padding-bottom: 0.3em;  /* 设置底部内边距 */
}

/* 设置二级标题的字体大小和底部边框 */
h2 {
    font-size: 2em;  /* 设置字体大小 */
    border-bottom: 1px solid #ddd;  /* 设置底部边框 */
    padding-bottom: 0.2em;  /* 设置底部内边距 */
}

/* 设置三级标题的字体大小 */
h3 {
    font-size: 1.75em;  /* 设置字体大小 */
}

/* 设置段落的行高和底部间距 */
p {
    line-height: 1.6;  /* 设置行高 */
    margin-bottom: 1.2em;  /* 设置底部间距 */
}

/* 设置链接的颜色和取消下划线 */
a {
    color: #0066cc;  /* 设置链接颜色 */
    text-decoration: none;  /* 取消下划线 */
}

/* 设置链接悬停时的下划线 */
a:hover {
    text-decoration: underline;  /* 悬停时显示下划线 */
}

/* 设置无序和有序列表的左边距和底部间距 */
ul, ol {
    margin-left: 20px;  /* 设置左边距 */
    margin-bottom: 1.2em;  /* 设置底部间距 */
}

/* 设置代码块的字体、背景颜色、内边距和圆角 */
code {
    font-family: Consolas, "Courier New", monospace;  /* 设置字体 */
    background-color: #f4f4f4;  /* 设置背景颜色 */
    padding: 2px 4px;  /* 设置内边距 */
    border-radius: 4px;  /* 设置圆角 */
}

/* 设置预格式化代码块的显示样式 */
pre code {
    display: block;  /* 代码块显示为块级元素 */
    padding: 10px;  /* 设置内边距 */
    background-color: #f4f4f4;  /* 设置背景颜色 */
    border: 1px solid #ddd;  /* 设置边框 */
    border-radius: 4px;  /* 设置圆角 */
    overflow-x: auto;  /* 超出时水平滚动 */
}

/* 设置引用块的样式 */
blockquote {
    border-left: 4px solid #ddd;  /* 设置左边框 */
    padding-left: 1em;  /* 设置左内边距 */
    color: #666;  /* 设置文本颜色 */
    margin: 1.2em 0;  /* 设置上下间距 */
    background-color: #f4f4f4;  /* 设置背景颜色 */
}

/* 设置表格的样式 */
table {
    width: 100%;  /* 设置表格宽度为100% */
    border-collapse: collapse;  /* 合并边框 */
    margin-bottom: 1.2em;  /* 设置底部间距 */
}

/* 设置表格、表头和单元格的边框和内边距 */
table, th, td {
    border: 1px solid #ddd;  /* 设置边框 */
    padding: 0.6em;  /* 设置内边距 */
}

/* 设置表头的背景颜色和对齐方式 */
th {
    background-color: #f2f2f2;  /* 设置背景颜色 */
    text-align: left;  /* 设置左对齐 */
}
""")

    # 创建临时Markdown文件并写入文档标题
    with open(temp_md_file, "w", encoding="utf-8") as f:
        f.write(f"% {title}\n\n")
        # 读取原始Markdown文件内容并写入临时Markdown文件
        with open(input_file, "r", encoding="utf-8") as original_md:
            previous_line = ""
            for line in original_md:
                # 每次遇到Markdown标题时在前面添加空行
                if line.strip().startswith("#"):
                    if previous_line.strip():
                        f.write("\n")
                    f.write("\n" + line.strip() + "\n\n")
                else:
                    f.write(line)
                previous_line = line

    # Pandoc命令，用于将Markdown转换为HTML
    command = [
        "pandoc",
        temp_md_file,  # 输入文件为临时Markdown文件
        "-o", output_file,  # 输出文件为指定的HTML文件
        "--self-contained",  # 生成包含所有资源的单个HTML文件
        "--resource-path", resource_path_str,  # 资源路径
        "-c", css_path  # 使用默认的CSS文件进行样式设置
    ]

    # 运行Pandoc命令
    result = subprocess.run(command, cwd=os.path.dirname(input_file), capture_output=True, text=True)

    # 检查命令执行结果，如果出错则打印错误信息
    if result.returncode != 0:
        print(f"Error converting {input_file} to {output_file}")
        print(result.stderr)

    # 删除临时Markdown文件
    os.remove(temp_md_file)


def convert_md_to_docx_with_toc_and_template(md_file_path, docx_file_path, template_file_path, title, version, date,
                                             left_header, right_header, statement, resource_paths, logo_path):
    """
    将Markdown文件转换为带有目录和模板的DOCX文件。

    参数:
        md_file_path (str): 输入的Markdown文件路径。
        docx_file_path (str): 输出的DOCX文件路径。
        template_file_path (str): DOCX模板文件路径。
        title (str): 文档标题。
        version (str): 版本号。
        date (str): 文档日期。
        left_header (str): 左页眉内容。
        right_header (str): 右页眉内容。
        statement (str): 可选声明。
        resource_paths (list): 资源文件路径列表。
        logo_path (str): Logo文件路径。
    """
    # 临时Markdown文件路径
    temp_md_file_path = 'temp.md'
    temp_docx_file_path = 'temp_output.docx'

    # 将资源路径列表转换为字符串，使用操作系统的路径分隔符
    resource_path_str = os.pathsep.join(resource_paths)

    # 创建一个临时的Markdown文件，用于存储转换过程中的中间数据
    with open(temp_md_file_path, "w", encoding="utf-8") as f:
        # 读取原始Markdown文件内容，并写入临时Markdown文件
        with open(md_file_path, "r", encoding="utf-8") as original_md:
            previous_line = ""
            for line in original_md:
                # 如果检测到图片标签，确保图片上下有空行
                if line.strip().startswith("![") and "](" in line:
                    if previous_line.strip():
                        f.write("\n")
                    f.write(line.strip() + "\n\n")
                elif line.strip().startswith('-'):
                    if previous_line.strip():
                        f.write("\n")
                    f.write(line.strip() + "\n\n")
                else:
                    f.write(line)
                previous_line = line

    # Pandoc命令
    pandoc_command = [
        'pandoc',
        temp_md_file_path,  # 输入文件为临时Markdown文件
        '-o', temp_docx_file_path,  # 输出文件为临时的DOCX文件
        '--toc',  # 启用目录
        '--toc-depth=3',  # 目录深度为3级
        '--reference-doc', template_file_path,  # 使用指定的DOCX模板
        '--resource-path', resource_path_str,  # 资源路径
        '--wrap=none',  # 不换行，保持原始格式
    ]

    # 运行Pandoc命令
    result = subprocess.run(pandoc_command, capture_output=True, text=True)

    # 检查命令执行结果
    if result.returncode == 0:
        print(f"Converted {md_file_path} to {docx_file_path} with template")

        # 创建一个新的DOCX文档并添加封面和目录
        final_doc = Document()
        add_cover_page(final_doc, title, version, date, statement)  # 添加封面页
        add_table_of_contents(final_doc)  # 添加目录页

        # 使用 Composer 合并生成的文档和 Pandoc 生成的文档内容
        composer = Composer(final_doc)
        doc_to_append = Document(temp_docx_file_path)
        composer.append(doc_to_append)

        # 添加页眉、页脚和首页页眉图片
        apply_headers_footers_to_sections(final_doc, left_header, right_header)
        add_header_image_to_first_page(final_doc, logo_path, right_text=right_header)

        # 保存最终文档
        final_doc.save(docx_file_path)

        # 删除临时文件
        os.remove(temp_docx_file_path)
        os.remove(temp_md_file_path)

    else:
        print(f"Error in conversion: {result.stderr}")